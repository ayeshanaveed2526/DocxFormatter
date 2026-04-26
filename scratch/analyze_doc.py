from docx import Document
import re

MCQ_OPTION_PATTERN    = re.compile(r'^\s*[a-dA-D][\.\)]\s+\S')
MCQ_QUESTION_PATTERN  = re.compile(r'^\s*Q?\d+[\.\)]\s+\S')
BULLET_PATTERN        = re.compile(r'^\s*[-•*]\s+\S')
NUMBERED_LIST_PATTERN = re.compile(r'^\s*\d+\.\s+[a-z]')
NUMBERED_SECTION_PATTERN = re.compile(r'^\d+\.\d+(\.\d+)?[\s\t]+\S')
MAIN_SECTION_PATTERN     = re.compile(r'^\d+\.\s+\S')

def classify_paragraph(para):
    text       = para.text.strip()
    style_name = para.style.name.lower() if para.style and para.style.name else ""
    if not text:
        return None
    if 'title' in style_name or 'heading 1' in style_name:
        return 'heading'
    if 'heading' in style_name:
        return 'subheading'
    if MCQ_OPTION_PATTERN.match(text):
        return 'option'
    if MCQ_QUESTION_PATTERN.match(text):
        return 'mcq'
    if BULLET_PATTERN.match(text) or NUMBERED_LIST_PATTERN.match(text):
        return 'paragraph'
    if MAIN_SECTION_PATTERN.match(text):
        return 'heading'
    if NUMBERED_SECTION_PATTERN.match(text):
        return 'subheading'
    words = text.split()
    if text.endswith(':') and len(words) <= 6:
        return 'heading' if len(words) <= 3 else 'subheading'
    if text.isupper() and len(words) <= 10:
        return 'heading'
    if len(words) <= 12 and not text.endswith(('.', ',', ';', ':', '?', '!')):
        return 'heading' if len(words) <= 5 else 'subheading'
    return 'paragraph'

doc = Document('backend/test_input.docx')
print("--- Paragraphs ---")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text: continue
    ptype = classify_paragraph(para)
    style = para.style.name if para.style else "None"
    print(f"[{ptype:10}] | Style: {style:15} | {text[:80]}")

print("\n--- Tables ---")
for i, table in enumerate(doc.tables):
    print(f"Table {i}:")
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                text = para.text.strip()
                if not text: continue
                ptype = classify_paragraph(para)
                print(f"  [{ptype:10}] | {text[:80]}")
