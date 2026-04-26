import re
import os
import json
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()

# --- Advanced Regex Patterns ---
MCQ_OPTION_PATTERN    = re.compile(r'^\s*([a-dA-D1-4]|[i-v]+)[\.\)]\s+\S')
MCQ_QUESTION_PATTERN  = re.compile(r'^\s*(Q|Question)?\s*\d+[\.\:\)]\s+\S', re.IGNORECASE)
BULLET_PATTERN        = re.compile(r'^\s*[-•*]\s+\S')
NUMBERED_LIST_PATTERN = re.compile(r'^\s*\d+\.\s+[a-z]')
NUMBERED_SECTION_PATTERN = re.compile(r'^\d+\.\d+(\.\d+)?[\s\t]+\S')
MAIN_SECTION_PATTERN     = re.compile(r'^\d+\.\s+[A-Z]')
URL_PATTERN           = re.compile(r'https?://\S+')
EMAIL_PATTERN         = re.compile(r'\S+@\S+\.\S+')

# --- Gemini Integration ---
def get_gemini_model():
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return None
    genai.configure(api_key=api_key)
    return genai.GenerativeModel('gemini-1.5-flash')

def extract_formatting_rules(instructions):
    """
    Converts natural language instructions into a structured JSON rule object.
    """
    if not instructions or not isinstance(instructions, str):
        return {}
    
    model = get_gemini_model()
    if not model:
        # Simple fallback parsing if no API key
        rules = {}
        if "size" in instructions.lower():
            match = re.search(r'size\s*(\d+)', instructions.lower())
            if match: rules["paragraph"] = {"size": int(match.group(1))}
        return rules

    prompt = f"""
    Convert these document formatting instructions into a JSON object.
    Instructions: "{instructions}"
    
    Format the JSON as follows:
    {{
        "heading": {{"size": int, "bold": bool}},
        "subheading": {{"size": int, "bold": bool}},
        "paragraph": {{"size": int, "bold": bool}},
        "mcq": {{"size": int, "bold": bool}},
        "option": {{"size": int, "bold": bool}},
        "margins": {{"top": float, "bottom": float, "left": float, "right": float}},
        "orientation": "portrait" | "landscape",
        "headerText": "string",
        "footerText": "string",
        "pageNumbers": "top" | "bottom" | "none"
    }}
    Only return the JSON.
    """
    try:
        response = model.generate_content(prompt)
        # Clean response text (remove markdown backticks if present)
        text = response.text.strip()
        if text.startswith("```json"):
            text = text[7:-3].strip()
        elif text.startswith("```"):
            text = text[3:-3].strip()
        return json.loads(text)
    except Exception as e:
        print(f"Gemini Error (Rule Extraction): {e}")
        return {}

def classify_paragraph(para, prev_type=None):
    text       = para.text.strip()
    style_name = para.style.name.lower() if para.style and para.style.name else ""
    
    if not text:
        return None
        
    # 1. Style-based detection (highest priority)
    if 'title' in style_name or 'heading 1' in style_name:
        return 'heading'
    if 'heading' in style_name:
        return 'subheading'
        
    # 2. MCQ detection
    if MCQ_QUESTION_PATTERN.match(text):
        return 'mcq'
    if MCQ_OPTION_PATTERN.match(text):
        return 'option'
    
    # 3. Contextual MCQ detection (if previous was MCQ/Option and this starts with a known option marker)
    if prev_type in ['mcq', 'option'] and re.match(r'^[a-dA-D1-4][\.\)]', text):
        return 'option'

    # 4. Pattern-based detection
    if MAIN_SECTION_PATTERN.match(text):
        return 'heading'
    if NUMBERED_SECTION_PATTERN.match(text):
        return 'subheading'
    if BULLET_PATTERN.match(text) or NUMBERED_LIST_PATTERN.match(text):
        return 'paragraph'
        
    # 5. Heuristics
    words = text.split()
    if not words: return None

    # Heading candidates: short, no punctuation, or uppercase
    if len(words) <= 12 and not text.endswith(('.', ',', ';', ':', '?', '!')):
        # Check if it was already bold in the original or is Title Case
        is_bold = any(run.bold for run in para.runs)
        is_title_case = all(w[0].isupper() for w in words if len(w) > 3 and w[0].isalpha())
        
        if is_bold or text.isupper() or (is_title_case and len(words) > 1):
            return 'heading' if len(words) <= 5 else 'subheading'
    
    if text.endswith(':') and len(words) <= 6:
        return 'heading' if len(words) <= 3 else 'subheading'
    
    # 6. Fallback for MCQ Options that might just be "A. text" without the pattern match
    if prev_type == 'mcq' and len(words) < 20:
        if re.match(r'^[A-D1-4][\.\)]', text):
            return 'option'

    return 'paragraph'

def set_table_format(table, table_rules):
    """
    Apply MS Word table formatting: alignment, borders, and width.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    # 1. Alignment
    align = table_rules.get("alignment", "left").lower()
    if align == "center":
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    elif align == "right":
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    else:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 2. Borders (Grid)
    if table_rules.get("borders", True):
        table.style = 'Table Grid'
    else:
        table.style = 'Normal Table'

def apply_formatting_to_para(para, ptype, rules):
    if not ptype or ptype not in rules:
        ptype = "paragraph"
        
    style_rule = rules.get(ptype, rules.get("paragraph", {"size": 12, "bold": False}))
    
    # --- Alignment ---
    align_str = style_rule.get("alignment", "left").lower()
    if align_str == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align_str == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align_str == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Typography ---
    for run in para.runs:
        if style_rule.get("size"):
            run.font.size = Pt(int(style_rule["size"]))
        if 'bold' in style_rule:
            run.bold = bool(style_rule["bold"])
        if style_rule.get("italic"):
            run.font.italic = True

def set_page_margins(doc, top, bottom, left, right):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

def set_page_orientation(doc, orientation):
    from docx.enum.section import WD_ORIENT
    for section in doc.sections:
        if orientation == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            if section.page_width < section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width

def add_page_number_field(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_header(doc, header_text, page_number_position):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        para = header.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_text:
            run = para.add_run(header_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if page_number_position == 'top':
            if header_text: para.add_run('  |  Page ')
            else: para.add_run('Page ')
            num_run = para.add_run()
            num_run.font.size = Pt(10)
            add_page_number_field(num_run)

def set_footer(doc, footer_text, page_number_position):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_text:
            run = para.add_run(footer_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if page_number_position == 'bottom':
            if footer_text: para.add_run('  |  Page ')
            else: para.add_run('Page ')
            num_run = para.add_run()
            num_run.font.size = Pt(10)
            add_page_number_field(num_run)
def process_document(input_path, output_path, rules_or_instructions):
    """
    Standardizes a document based on structured rules or natural language instructions.
    """
    if isinstance(rules_or_instructions, str):
        rules = extract_formatting_rules(rules_or_instructions)
    else:
        rules = rules_or_instructions

    if not rules:
        rules = {
            "paragraph": {"size": 12, "bold": False, "alignment": "left"},
            "heading": {"size": 16, "bold": True, "alignment": "left"},
            "subheading": {"size": 14, "bold": True, "alignment": "left"},
            "table": {"alignment": "center", "borders": True}
        }

    doc = Document(input_path)
    
    # 2. Page Setup
    margins = rules.get("margins", {})
    if margins:
        set_page_margins(
            doc,
            top    = float(margins.get("top",    2.54)),
            bottom = float(margins.get("bottom", 2.54)),
            left   = float(margins.get("left",   2.54)),
            right  = float(margins.get("right",  2.54)),
        )
    set_page_orientation(doc, rules.get("orientation", "portrait"))
    
    set_header(doc, rules.get("headerText", ""), rules.get("pageNumbers", "none"))
    set_footer(doc, rules.get("footerText", ""), rules.get("pageNumbers", "none"))

    # 3. Process Content
    table_rules = rules.get("table", {"alignment": "center", "borders": True})

    def process_paras(paragraphs):
        prev_type = None
        for para in paragraphs:
            if not para.text.strip():
                # Center images if they are the only thing in the paragraph
                if 'w:drawing' in para._p.xml or 'w:pict' in para._p.xml:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
            ptype = classify_paragraph(para, prev_type)
            if ptype:
                apply_formatting_to_para(para, ptype, rules)
                prev_type = ptype

    process_paras(doc.paragraphs)
    
    for table in doc.tables:
        set_table_format(table, table_rules)
        for row in table.rows:
            for cell in row.cells:
                process_paras(cell.paragraphs)

    doc.save(output_path)
    print(f"Saved formatted document to: {output_path}")

# Compatibility function for test_gemini.py
def process_text_batch(texts, instructions):
    """
    Mocking AI classification for batch processing if needed by test scripts.
    """
    rules = extract_formatting_rules(instructions)
    results = []
    prev_type = None
    for text in texts:
        # Create a dummy paragraph for classify_paragraph
        class DummyPara:
            def __init__(self, t):
                self.text = t
                self.style = None
                self.runs = []
        
        ptype = classify_paragraph(DummyPara(text), prev_type)
        results.append({"text": text, "type": ptype or "paragraph"})
        prev_type = ptype
    return results
